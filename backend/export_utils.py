# backend/export_utils.py

from docx import Document
from docx.shared import Pt
from typing import List
from .models import Answer

def create_docx(answers: List[Answer]) -> bytes:
    """Generate a Word document (.docx) from a list of Answer objects."""
    doc = Document()
    doc.add_heading('Questionnaire Answers – PulseCRM', level=0)
    
    for ans in answers:
        # Question
        p = doc.add_paragraph()
        p.add_run(f"Q{ans.question.number}: ").bold = True
        p.add_run(ans.question.text)
        
        # Answer
        p = doc.add_paragraph()
        p.add_run("Answer: ").bold = True
        if ans.not_found:
            p.add_run(ans.text).italic = True
        else:
            p.add_run(ans.text)
        
        # Citations
        if ans.citations:
            p = doc.add_paragraph()
            run = p.add_run("Citations:")
            run.font.size = Pt(10)
            citations_text = "\n".join(ans.citations)
            run = p.add_run("\n" + citations_text)
            run.font.size = Pt(9)
            # Indent citations
            p.paragraph_format.left_indent = Pt(20)
    
    # Save to in-memory bytes
    from io import BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio